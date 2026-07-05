import re
from io import BytesIO
from pathlib import Path
from zipfile import BadZipFile

from docx import Document
from openpyxl import load_workbook
from pypdf import PdfReader


def extract_text_from_bytes(source_name: str, data: bytes) -> str:
    suffix = Path(source_name.split("?", 1)[0]).suffix.lower()
    if suffix == ".pdf":
        return _pdf_text(data)
    if suffix == ".docx":
        return _docx_text(data)
    if suffix == ".xlsx":
        return _xlsx_text(data)
    if suffix in {".xml", ".html", ".htm", ".txt"}:
        return data.decode("utf-8", errors="replace")
    return data.decode("utf-8", errors="replace")


def _pdf_text(data: bytes) -> str:
    reader = PdfReader(BytesIO(data))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _docx_text(data: bytes) -> str:
    doc = Document(BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            paragraphs.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(paragraphs)


def _xlsx_text(data: bytes) -> str:
    try:
        wb = load_workbook(BytesIO(data), data_only=True, read_only=True)
    except BadZipFile:
        return data.decode("utf-8", errors="replace")
    lines: list[str] = []
    for sheet in wb.worksheets:
        lines.append(f"# {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            values = [str(cell).strip() for cell in row if cell not in (None, "")]
            if values:
                lines.append(" | ".join(values))
    return "\n".join(lines)


def strip_html(value: str) -> str:
    return re.sub(r"\s+", " ", re.sub(r"<[^>]+>", " ", value)).strip()

